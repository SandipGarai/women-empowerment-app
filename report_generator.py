"""
Report generator for the Women Empowerment analysis GUI.
Creates Word (.docx) and PDF (.pdf) reports from the analysis outputs.
"""

from pathlib import Path
import re

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from fpdf.enums import XPos, YPos


PROJECT_DIR = Path(__file__).resolve().parent
DEFAULT_LOGO = None


def clean_text(value):
    """Convert any value to safe text for reports."""
    if value is None:
        return ""
    text = str(value)
    # Remove characters that break Word/XML or PDF output
    text = text.replace("\r", " ")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", text)
    return text.strip()


def clean_text_pdf(value):
    """Convert text to ASCII-safe for basic PDF fonts."""
    text = clean_text(value)
    # Replace common Unicode characters with ASCII equivalents
    replacements = {
        "→": "->",
        "←": "<-",
        "–": "-",
        "—": "-",
        "“": '"',
        "”": '"',
        "‘": "'",
        "’": "'",
        "•": "*",
        "●": "*",
        "\u2010": "-",
        "\u2011": "-",
        "\u2012": "-",
        "\u2013": "-",
        "\u2014": "-",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Remove remaining non-latin1 characters
    text = text.encode("latin-1", "ignore").decode("latin-1")
    return text


def add_heading(doc, text, level=1):
    """Add a heading to Word document."""
    heading = doc.add_heading(level=level)
    run = heading.add_run(clean_text(text))
    run.bold = True
    return heading


def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph to Word document."""
    p = doc.add_paragraph()
    run = p.add_run(clean_text(text))
    run.bold = bold
    run.italic = italic
    return p


def generate_word_report(
    output_path,
    title="Women Empowerment Analysis Report",
    author="",
    objective_mapping_path=None,
    all_results_path=None,
    question_summary_path=None,
    validation_status_path=None,
    include_tables=True,
):
    """Generate a Word report from analysis outputs."""

    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(clean_text(title))
    run.bold = True
    run.font.size = Pt(18)

    if author:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run(f"Author: {clean_text(author)}")
        run.italic = True

    doc.add_paragraph()

    # Summary
    add_heading(doc, "Executive Summary", level=1)
    add_paragraph(
        doc,
        "This report presents findings from the survey analysis on democratic decentralization "
        "and women empowerment among Oraon tribes in Mandar block, Ranchi district.",
    )

    # Validation status
    if validation_status_path and Path(validation_status_path).exists():
        add_heading(doc, "Validation Status", level=1)
        status_df = pd.read_csv(validation_status_path)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Check"
        hdr_cells[1].text = "Status"
        hdr_cells[2].text = "Note"
        for _, row in status_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = clean_text(row["check"])
            row_cells[1].text = clean_text(row["status"])
            row_cells[2].text = clean_text(row["note"])
        doc.add_paragraph()

    # Findings by objective
    if objective_mapping_path and Path(objective_mapping_path).exists():
        add_heading(doc, "Findings by Synopsis Objective", level=1)
        mapping_df = pd.read_excel(objective_mapping_path)
        numeric_df = mapping_df[mapping_df["top_percent"].notna()].copy()
        numeric_df["top_percent"] = pd.to_numeric(numeric_df["top_percent"], errors="coerce")

        for obj in mapping_df["objective"].unique():
            add_heading(doc, clean_text(obj), level=2)
            sub = numeric_df[numeric_df["objective"] == obj]
            if len(sub) == 0:
                add_paragraph(doc, "No numeric findings available for this objective yet.")
                continue
            for _, row in sub.iterrows():
                pct = row["top_percent"]
                pct_str = f"{pct:.1f}%" if pd.notna(pct) else ""
                line = (
                    f"{row['respondent_group']} Q{row['question_no']}: "
                    f"{row['top_answer']} ({row['top_count']}/{row['denominator_n']}, {pct_str}). "
                    f"Question: {row['question']}"
                )
                add_paragraph(doc, line)
            doc.add_paragraph()

    # Detailed results tables
    if include_tables and all_results_path and Path(all_results_path).exists():
        add_heading(doc, "Detailed Results", level=1)
        results_df = pd.read_csv(all_results_path)
        themes = sorted(results_df["theme"].unique())
        for theme in themes:
            add_heading(doc, clean_text(theme), level=2)
            theme_df = results_df[results_df["theme"] == theme]
            # Limit rows per theme to keep report readable
            sample_df = theme_df.head(30)
            table = doc.add_table(rows=1, cols=5)
            table.style = "Light Grid Accent 1"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Group"
            hdr_cells[1].text = "Q No"
            hdr_cells[2].text = "Question"
            hdr_cells[3].text = "Answer"
            hdr_cells[4].text = "Percent"
            for _, row in sample_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = clean_text(row["respondent_group"])
                row_cells[1].text = clean_text(row["question_no"])
                row_cells[2].text = clean_text(row["question"])[:120]
                row_cells[3].text = clean_text(row["answer"])[:120]
                pct = row["percent_of_respondents"]
                row_cells[4].text = f"{pct:.1f}%" if pd.notna(pct) else ""
            doc.add_paragraph()

    # Question-level summary
    if question_summary_path and Path(question_summary_path).exists():
        add_heading(doc, "Question-Level Summary", level=1)
        qsum_df = pd.read_csv(question_summary_path)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Theme"
        hdr_cells[1].text = "Group"
        hdr_cells[2].text = "Q No"
        hdr_cells[3].text = "Top Answer"
        hdr_cells[4].text = "Count"
        hdr_cells[5].text = "Percent"
        for _, row in qsum_df.head(60).iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = clean_text(row["theme"])
            row_cells[1].text = clean_text(row["respondent_group"])
            row_cells[2].text = clean_text(row["question_no"])
            row_cells[3].text = clean_text(row["top_answer"])[:100]
            row_cells[4].text = clean_text(row["top_count"])
            pct = row["top_percent"]
            row_cells[5].text = f"{pct:.1f}%" if pd.notna(pct) else ""
        doc.add_paragraph()

    doc.save(output_path)
    return output_path


# Locate DejaVu fonts: bundled fonts dir first, then system paths
_BUNDLED_FONT_DIR = PROJECT_DIR / "fonts"
_SYSTEM_FONT_DIRS = [
    Path("/usr/share/fonts/truetype/dejavu"),  # Linux
    Path("/usr/local/share/fonts/truetype/dejavu"),  # macOS / custom Linux
    Path.home() / "AppData" / "Local" / "Microsoft" / "Windows" / "Fonts",  # Windows user
    Path("C:/Windows/Fonts"),  # Windows system
]


def _find_font(filename):
    """Return the path to a DejaVu TTF font file."""
    candidates = [_BUNDLED_FONT_DIR / filename]
    candidates += [d / filename for d in _SYSTEM_FONT_DIRS]
    for path in candidates:
        if path.exists():
            return str(path)
    raise FileNotFoundError(
        f"DejaVu font file '{filename}' not found. "
        f"Please copy the DejaVu fonts into {_BUNDLED_FONT_DIR} or install them on your system."
    )


class PDFReport(FPDF):
    """Simple PDF report class using Unicode-safe font."""

    def __init__(self):
        super().__init__()
        self.add_font("DejaVu", "", _find_font("DejaVuSans.ttf"))
        self.add_font("DejaVu", "B", _find_font("DejaVuSans-Bold.ttf"))
        self.add_font("DejaVu", "I", _find_font("DejaVuSans-Oblique.ttf"))

    def header(self):
        self.set_font("DejaVu", "B", 12)
        self.cell(0, 10, "Women Empowerment Analysis Report", border=0, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
        self.ln(2)

    def footer(self):
        self.set_y(-15)
        self.set_font("DejaVu", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")

    def chapter_title(self, title):
        self.set_font("DejaVu", "B", 14)
        self.set_fill_color(230, 230, 230)
        self.cell(0, 10, clean_text_pdf(title)[:100], new_x=XPos.LMARGIN, new_y=YPos.NEXT, fill=True)
        self.ln(2)

    def chapter_subtitle(self, subtitle):
        self.set_font("DejaVu", "B", 11)
        self.cell(0, 8, clean_text_pdf(subtitle)[:120], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.ln(1)

    def chapter_body(self, body):
        self.set_font("DejaVu", "", 10)
        text = clean_text_pdf(body)
        self.multi_cell(0, 6, text)
        self.ln(2)


def generate_pdf_report(
    output_path,
    title="Women Empowerment Analysis Report",
    objective_mapping_path=None,
):
    """Generate a simple PDF report."""
    pdf = PDFReport()
    pdf.add_page()
    pdf.set_font("DejaVu", "B", 16)
    pdf.cell(0, 10, clean_text_pdf(title), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(5)

    pdf.chapter_title("Executive Summary")
    pdf.chapter_body(
        "This report presents findings from the survey analysis on democratic decentralization "
        "and women empowerment among Oraon tribes in Mandar block, Ranchi district."
    )

    if objective_mapping_path and Path(objective_mapping_path).exists():
        pdf.chapter_title("Findings by Synopsis Objective")
        mapping_df = pd.read_excel(objective_mapping_path)
        numeric_df = mapping_df[mapping_df["top_percent"].notna()].copy()
        numeric_df["top_percent"] = pd.to_numeric(numeric_df["top_percent"], errors="coerce")

        for obj in mapping_df["objective"].unique():
            pdf.chapter_subtitle(clean_text(obj))
            sub = numeric_df[numeric_df["objective"] == obj]
            if len(sub) == 0:
                pdf.chapter_body("No numeric findings available for this objective yet.")
                continue
            for _, row in sub.iterrows():
                pct = row["top_percent"]
                pct_str = f"{pct:.1f}%" if pd.notna(pct) else ""
                line = (
                    f"{row['respondent_group']} Q{row['question_no']}: "
                    f"{row['top_answer']} ({row['top_count']}/{row['denominator_n']}, {pct_str})"
                )
                pdf.chapter_body(line)

    pdf.output(output_path)
    return output_path


if __name__ == "__main__":
    # Test run
    out_word = PROJECT_DIR / "outputs" / "gui_report_test.docx"
    out_pdf = PROJECT_DIR / "outputs" / "gui_report_test.pdf"
    generate_word_report(
        out_word,
        objective_mapping_path=PROJECT_DIR / "outputs" / "synopsis_mapping" / "objectives_with_top_findings.xlsx",
        all_results_path=PROJECT_DIR / "outputs" / "descriptive_results" / "all_descriptive_results.csv",
        question_summary_path=PROJECT_DIR / "outputs" / "descriptive_results" / "question_level_summary.csv",
        validation_status_path=PROJECT_DIR / "outputs" / "validation" / "validation_status_summary.csv",
    )
    generate_pdf_report(
        out_pdf,
        objective_mapping_path=PROJECT_DIR / "outputs" / "synopsis_mapping" / "objectives_with_top_findings.xlsx",
    )
    print(f"Test Word report: {out_word}")
    print(f"Test PDF report: {out_pdf}")
